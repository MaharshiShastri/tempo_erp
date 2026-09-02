from database.repository import EDBR

from pathlib import Path
from copy import deepcopy
import json
import os
import re
from difflib import SequenceMatcher

from groq import Groq
from decimal import Decimal, ROUND_HALF_UP
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
)
from docx.enum.text import WD_COLOR_INDEX

# ============================================================
# Paths / Configuration
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent

TEMPLATE_PATH = (
    BASE_DIR
    / "services"
    / "template_document.docx"
)

PRICE_LIST_PATH = (
    BASE_DIR
    / "services"
    / "Ex Works Price List 2026-27.docx"
)

OUTPUT_DIR = (
    BASE_DIR
    / "generated_quotes"
)


# ============================================================
# Groq Configuration
# ============================================================

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

if not GROQ_API_KEY:
    raise RuntimeError(
        "GROQ_API_KEY environment variable is not set."
    )

groq_client = Groq(
    api_key=GROQ_API_KEY
)

GROQ_MODEL = "openai/gpt-oss-20b"


# ============================================================
# General Helpers
# ============================================================

def safe_filename(value: str) -> str:
    """
    Convert a value into a filesystem-safe filename component.
    """

    value = str(value or "").strip()

    value = re.sub(
        r'[<>:"/\\|?*]',
        "_",
        value,
    )

    value = value.rstrip(" .")

    return value

def money(value):
    if value is None:
        return Decimal("0.00")

    return Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

def calculate_quotation_financials(request):
    base = money(request.base_model_price)

    packing = money(request.packing_amount) if request.packing_mode == "ACTUAL" else Decimal("0.00")

    freight = money(request.freight_amount) if request.freight_mode == "ACTUAL" else Decimal("0.00")

    taxable_value = money(base + packing + freight)

    tax_amount = money(taxable_value * money(request.tax_rate) / Decimal("100"))

    grand_total = money(taxable_value + tax_amount)

    return {
        "base_model_price": base,
        "packing_amount": packing,
        "freight_amount": freight,
        "taxable_value": taxable_value,
        "tax_rate": money(request.tax_rate),
        "tax_amount": tax_amount,
        "grand_total": grand_total,
    }

def normalize_text(value: str) -> str:
    """
    Normalize ordinary document text.
    """

    return " ".join(
        str(value or "")
        .strip()
        .upper()
        .split()
    )


def normalize_heading(value: str) -> str:
    """
    Normalize product-group headings.
    """

    value = str(value or "").upper().strip()

    value = re.sub(
        r"\s*[:\-]+\s*$",
        "",
        value,
    )

    value = re.sub(
        r"\s*-\s*",
        "-",
        value,
    )

    return " ".join(
        value.split()
    )


def normalize_item_code(value: str) -> str:
    """
    Normalize an item/model code.

    Examples:

        TI-90A
        TI 90A
        ti_90a
        TI/90A

    become:

        TI 90A
    """

    value = str(value or "").upper().strip()

    value = re.sub(
        r"[^A-Z0-9]+",
        " ",
        value,
    )

    return " ".join(
        value.split()
    )


def compact_item_code(value: str) -> str:
    """
    Remove formatting characters from an item code.

    Examples:

        TI-90A  -> TI90A
        TI 90A  -> TI90A
        TI_90A  -> TI90A
    """

    return re.sub(
        r"[^A-Z0-9]",
        "",
        str(value or "").upper(),
    )


def item_tokens(value: str) -> set:
    """
    Return normalized item-code tokens.
    """

    normalized = normalize_item_code(
        value
    )

    if not normalized:
        return set()

    return set(
        normalized.split()
    )


def item_match_score(
    requested: str,
    candidate: str,
) -> int:
    """
    Score similarity between two item/model codes.

    Exact matches always receive the highest score.

    The function deliberately gives priority to the actual
    model identifier rather than loose string similarity.
    """

    requested_normalized = normalize_item_code(
        requested
    )

    candidate_normalized = normalize_item_code(
        candidate
    )

    if (
        not requested_normalized
        or not candidate_normalized
    ):
        return 0

    # --------------------------------------------------------
    # Exact normalized match
    # --------------------------------------------------------

    if (
        requested_normalized
        == candidate_normalized
    ):
        return 1000

    # --------------------------------------------------------
    # Compact exact match
    # --------------------------------------------------------

    requested_compact = compact_item_code(
        requested
    )

    candidate_compact = compact_item_code(
        candidate
    )

    if (
        requested_compact
        and requested_compact
        == candidate_compact
    ):
        return 980

    requested_tokens = item_tokens(
        requested
    )

    candidate_tokens = item_tokens(
        candidate
    )

    if (
        not requested_tokens
        or not candidate_tokens
    ):
        return 0

    common_tokens = (
        requested_tokens
        & candidate_tokens
    )

    if not common_tokens:
        return 0

    overlap = (
        len(common_tokens)
        / max(
            len(requested_tokens),
            len(candidate_tokens),
        )
    )

    # --------------------------------------------------------
    # Candidate is effectively a base model
    #
    # Requested:
    #   TI 715 WIC SPL
    #
    # Candidate:
    #   TI 715 WIC
    # --------------------------------------------------------

    if candidate_tokens.issubset(
        requested_tokens
    ):
        return 900 + int(
            overlap * 50
        )

    # --------------------------------------------------------
    # Requested code is contained in candidate
    # --------------------------------------------------------

    if requested_tokens.issubset(
        candidate_tokens
    ):
        return 850 + int(
            overlap * 50
        )

    # --------------------------------------------------------
    # Similarity fallback
    # --------------------------------------------------------

    similarity = SequenceMatcher(
        None,
        requested_normalized,
        candidate_normalized,
    ).ratio()

    return int(
        similarity * 100
    )


# ============================================================
# Item-Code Matching
# ============================================================

def cell_contains_item_code(
    cell,
    requested_item_code,
) -> bool:
    """
    Determine whether a Word cell contains the requested
    product/model code.

    This is intentionally more robust than simple substring
    matching.

    Supported examples:

        requested = TI-90A

        matches:
            TI-90A
            TI 90A
            TI90A
            TI-90A WIC
            TI 90A WIC
            Model: TI-90A

        does NOT match:

            TI-90AB
            TI-90B
            TI-901
            TI-90
    """

    requested_normalized = normalize_item_code(
        requested_item_code
    )

    if not requested_normalized:
        return False

    cell_text = normalize_text(
        cell.text
    )

    if not cell_text:
        return False

    cell_normalized = normalize_item_code(
        cell_text
    )

    # --------------------------------------------------------
    # Exact normalized cell
    # --------------------------------------------------------

    if (
        cell_normalized
        == requested_normalized
    ):
        return True

    # --------------------------------------------------------
    # Token-based matching
    #
    # This allows:
    #
    # TI 90A WIC
    #
    # to match:
    #
    # TI 90A
    # --------------------------------------------------------

    requested_tokens = item_tokens(
        requested_item_code
    )

    cell_tokens = item_tokens(
        cell_text
    )

    if (
        requested_tokens
        and requested_tokens.issubset(
            cell_tokens
        )
    ):
        return True

    # --------------------------------------------------------
    # Compact model matching
    #
    # TI-90A -> TI90A
    #
    # We require a proper alphanumeric boundary so:
    #
    # TI90A
    #
    # matches:
    #
    # TI90A WIC
    #
    # but:
    #
    # TI90AB
    #
    # does not.
    # --------------------------------------------------------

    requested_compact = compact_item_code(
        requested_item_code
    )

    cell_compact = compact_item_code(
        cell_text
    )

    if (
        requested_compact
        and cell_compact
    ):

        pattern = (
            rf"(?<![A-Z0-9])"
            rf"{re.escape(requested_compact)}"
            rf"(?![A-Z0-9])"
        )

        if re.search(
            pattern,
            cell_compact,
        ):
            return True

    # --------------------------------------------------------
    # Search original text using normalized code variants.
    #
    # This handles cases where Word exposes formatting
    # differently from the expected text representation.
    # --------------------------------------------------------

    original_upper = str(
        cell.text or ""
    ).upper()

    compact_original = re.sub(
        r"[^A-Z0-9]",
        "",
        original_upper,
    )

    if (
        requested_compact
        and requested_compact
        == compact_original
    ):
        return True

    # Requested model can be followed by suffixes.
    if requested_compact:

        if compact_original.startswith(
            requested_compact
        ):

            remainder = (
                compact_original[
                    len(requested_compact):
                ]
            )

            # Don't allow another alphanumeric
            # model character immediately after it.
            if not remainder:
                return True

            if not remainder[0].isalnum():
                return True

    return False


def row_contains_item_code(
    row,
    requested_item_code,
) -> bool:
    """
    Check every cell in a Word table row.

    This is important because price-list documents often
    distribute model information across multiple cells.
    """

    for cell in row.cells:

        if cell_contains_item_code(
            cell,
            requested_item_code,
        ):
            return True

    return False


def row_text(
    row,
) -> str:
    """
    Return normalized text from all cells in a row.
    """

    return " | ".join(
        normalize_text(
            cell.text
        )
        for cell in row.cells
    )


def row_match_score(
    row,
    requested_item_code,
) -> int:
    """
    Score an entire table row against the requested item code.

    The best cell match determines the row score.
    """

    best_score = 0

    for cell in row.cells:

        cell_text = normalize_text(
            cell.text
        )

        if not cell_text:
            continue

        # Direct containment is strongest.
        if cell_contains_item_code(
            cell,
            requested_item_code,
        ):
            best_score = max(
                best_score,
                950,
            )

        # Score the actual cell text.
        best_score = max(
            best_score,
            item_match_score(
                requested_item_code,
                cell_text,
            ),
        )

    return best_score


# ============================================================
# Document Traversal
# ============================================================

def iter_block_items(parent):
    """
    Iterate through paragraphs and tables in document order.

    This preserves the original Word document ordering.
    """

    if isinstance(
        parent,
        _Document,
    ):

        parent_element = (
            parent.element.body
        )

    else:

        parent_element = (
            parent._tc
        )

    for child in (
        parent_element.iterchildren()
    ):

        if child.tag.endswith("}p"):

            yield Paragraph(
                child,
                parent,
            )

        elif child.tag.endswith("}tbl"):

            yield Table(
                child,
                parent,
            )


# ============================================================
# Product Group Discovery
# ============================================================

def find_product_group_bounds(
    source_doc,
    product_group,
    available_groups,
):
    """
    Return:

        (group_start, group_end, blocks)

    where group_start and group_end are indexes into the
    document's ordered blocks.
    """

    target_group = normalize_heading(
        product_group
    )

    normalized_groups = {
        normalize_heading(group)
        for group in (
            available_groups or []
        )
        if group
    }

    blocks = list(
        iter_block_items(source_doc)
    )

    group_start = None
    group_end = len(blocks)

    for index, block in enumerate(
        blocks
    ):

        if not isinstance(
            block,
            Paragraph,
        ):
            continue

        text = normalize_heading(block.text)

        if text == target_group:

            group_start = index

            continue

        if (
            group_start is not None
            and index > group_start
            and text in normalized_groups
            and text != target_group
        ):

            group_end = index

            break

    if group_start is None:

        raise ValueError(
            f"Product group "
            f"'{product_group}' not found "
            f"in the price-list document."
        )

    return (
        group_start,
        group_end,
        blocks,
    )


def find_product_group_tables(
    source_doc,
    product_group,
    available_groups,
):
    """
    Return all tables belonging to the requested product group.
    """

    (
        group_start,
        group_end,
        blocks,
    ) = find_product_group_bounds(
        source_doc=source_doc,
        product_group=product_group,
        available_groups=available_groups,
    )

    tables = []

    for block in blocks[
        group_start + 1 : group_end
    ]:

        if isinstance(
            block,
            Table,
        ):

            tables.append(block)

    return tables


# ============================================================
# Word Table Conversion
# ============================================================

def word_table_to_data(
    table,
):
    """
    Convert a Word table into a JSON-safe structure.

    The original table is never modified.
    """

    rows = []

    for row_index, row in enumerate(
        table.rows
    ):

        values = []

        for cell in row.cells:

            values.append(
                normalize_text(
                    cell.text
                )
            )

        rows.append(
            {
                "row_index": row_index,
                "values": values,
            }
        )

    return rows


# ============================================================
# Deterministic Table Identification
# ============================================================

def find_table_containing_item(
    tables,
    item_code,
):
    """
    Find the price-list table containing the requested model.

    Important:

    The previous implementation only did:

        requested in normalized_row

    which is unreliable for Word documents.

    This implementation:

        1. checks every cell
        2. scores every row
        3. selects the strongest table
        4. only fails when there is genuinely no deterministic
           indication of the requested model

    Returns:

        Word Table
    """

    if not tables:

        raise ValueError(
            "No price-list tables were supplied "
            "for item lookup."
        )

    requested = normalize_item_code(
        item_code
    )

    if not requested:

        raise ValueError(
            "Item code is required for "
            "price-list lookup."
        )

    # --------------------------------------------------------
    # First pass: direct cell matching
    # --------------------------------------------------------

    direct_matches = []

    for table_index, table in enumerate(
        tables
    ):

        for row_index, row in enumerate(
            table.rows
        ):

            # Skip header.
            if row_index == 0:
                continue

            if row_contains_item_code(
                row,
                requested,
            ):

                score = row_match_score(
                    row,
                    requested,
                )

                direct_matches.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "score": score,
                        "row_text": row_text(row),
                    }
                )

    if direct_matches:

        direct_matches.sort(
            key=lambda item: (
                item["score"],
                -item["table_index"],
            ),
            reverse=True,
        )

        best = direct_matches[0]

        return tables[
            best["table_index"]
        ]

    # --------------------------------------------------------
    # Second pass: deterministic fuzzy scoring
    # --------------------------------------------------------

    candidates = []

    for table_index, table in enumerate(
        tables
    ):

        for row_index, row in enumerate(
            table.rows
        ):

            if row_index == 0:
                continue

            score = row_match_score(
                row,
                requested,
            )

            if score > 0:

                candidates.append(
                    {
                        "table_index": table_index,
                        "row_index": row_index,
                        "score": score,
                        "row_text": row_text(row),
                    }
                )

    if candidates:

        candidates.sort(
            key=lambda item: (
                item["score"],
                -item["table_index"],
            ),
            reverse=True,
        )

        best = candidates[0]

        # Only accept a meaningful deterministic match.
        if best["score"] >= 70:

            return tables[
                best["table_index"]
            ]

    # --------------------------------------------------------
    # No deterministic table match.
    #
    # If only one table exists, it is necessarily the relevant
    # table inside this product group.
    # --------------------------------------------------------

    if len(tables) == 1:

        return tables[0]

    # --------------------------------------------------------
    # Multiple tables remain ambiguous.
    #
    # Use Groq to identify the table, but NOT the product data.
    # --------------------------------------------------------

    return find_table_with_groq(
        tables=tables,
        requested_item_code=item_code,
    )


# ============================================================
# Groq Table Selection
# ============================================================

def find_table_with_groq(
    tables,
    requested_item_code,
):
    """
    Ask Groq which candidate table contains the requested item.

    Groq receives only table data and returns a table index.

    It does not modify the document.
    """

    table_data = []

    for table_index, table in enumerate(
        tables
    ):

        table_data.append(
            {
                "table_index": table_index,
                "rows": word_table_to_data(
                    table
                ),
            }
        )

    prompt = f"""
You are selecting one table from a quotation price-list
document.

Requested item code:

{requested_item_code}

Candidate tables:

{json.dumps(
    table_data,
    ensure_ascii=False,
    indent=2,
)}

Your task:

1. Ignore header rows.
2. Find the candidate table containing the requested
   product/model.
3. Match formatting variants intelligently.

For example:

TI-90A
TI 90A
TI90A
TI-90A WIC

may represent the same base model.

However:

TI-90AB
TI-90B
TI-901
TI-90

must NOT be treated as TI-90A.

The actual model identifier must correspond.

Do NOT invent a table.

Do NOT invent product information.

Return ONLY valid JSON:

{{
    "selected_table_index": 0,
    "confidence": 0.99
}}

If no reliable table exists:

{{
    "selected_table_index": null,
    "confidence": 0.0
}}
"""

    response = groq_client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a deterministic "
                    "price-list table selector. "
                    "Return only valid JSON."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        temperature=0,
        response_format={
            "type": "json_object",
        },
    )

    content = (
        response
        .choices[0]
        .message
        .content
    )

    if not content:

        raise ValueError(
            "Groq returned an empty response "
            "while selecting the price-list table."
        )

    try:

        result = json.loads(
            content
        )

    except json.JSONDecodeError as exc:

        raise ValueError(
            "Groq returned invalid JSON while "
            "selecting the price-list table."
        ) from exc

    selected_table_index = result.get(
        "selected_table_index"
    )

    try:

        confidence = float(
            result.get(
                "confidence",
                0,
            )
        )

    except (
        TypeError,
        ValueError,
    ):

        confidence = 0.0

    if selected_table_index is None:

        raise ValueError(
            f"Could not determine which "
            f"price-list table contains item "
            f"'{requested_item_code}'."
        )

    try:

        selected_table_index = int(
            selected_table_index
        )

    except (
        TypeError,
        ValueError,
    ) as exc:

        raise ValueError(
            "Groq returned an invalid "
            "price-list table index."
        ) from exc

    if (
        selected_table_index < 0
        or selected_table_index >= len(tables)
    ):

        raise ValueError(
            "Groq returned an out-of-range "
            "price-list table index."
        )

    if confidence < 0.90:

        raise ValueError(
            f"Groq table-match confidence too low "
            f"for '{requested_item_code}': "
            f"{confidence}"
        )

    return tables[
        selected_table_index
    ]


# ============================================================
# Deterministic Product Row Identification
# ============================================================

def find_table_row_deterministically(
    table,
    requested_item_code,
):
    """
    Find the requested row without using Groq.

    Returns:

        row_index

    or:

        None
    """

    exact_matches = []

    scored_matches = []

    for row_index, row in enumerate(
        table.rows
    ):

        if row_index == 0:
            continue

        score = row_match_score(
            row,
            requested_item_code,
        )

        if score >= 950:

            exact_matches.append(
                row_index
            )

        elif score >= 70:

            scored_matches.append(
                (
                    score,
                    row_index,
                )
            )

    # --------------------------------------------------------
    # Exact/direct match
    # --------------------------------------------------------

    if exact_matches:

        return exact_matches[0]

    # --------------------------------------------------------
    # Strong fuzzy match
    # --------------------------------------------------------

    if scored_matches:

        scored_matches.sort(
            key=lambda item: item[0],
            reverse=True,
        )

        best_score, best_row = (
            scored_matches[0]
        )

        # Don't accept weak fuzzy matching.
        if best_score >= 80:

            return best_row

    return None


# ============================================================
# Groq Product Row Selection
# ============================================================

def find_table_row_with_groq(
    table,
    requested_item_code,
):
    """
    Select the requested product row from a price-list table.

    Deterministic matching is always attempted first.

    Groq is only used when deterministic matching cannot
    identify the row confidently.
    """

    # --------------------------------------------------------
    # Deterministic first
    # --------------------------------------------------------

    deterministic_row = (
        find_table_row_deterministically(
            table=table,
            requested_item_code=requested_item_code,
        )
    )

    if deterministic_row is not None:

        return deterministic_row

    # --------------------------------------------------------
    # Groq fallback
    # --------------------------------------------------------

    table_data = word_table_to_data(
        table
    )

    prompt = f"""
You are selecting ONE product row from a quotation
price-list table.

Requested item code:

{requested_item_code}

Table:

{json.dumps(
    table_data,
    ensure_ascii=False,
    indent=2,
)}

Rules:

1. Ignore row 0 because it is the header.
2. Select only one row.
3. The requested model must actually correspond to the
   model identifier in the row.
4. Formatting differences are acceptable.

Examples:

Requested:
TI-90A

Valid equivalents:
TI-90A
TI 90A
TI90A
TI-90A WIC

Invalid:

TI-90AB
TI-90B
TI-901
TI-90

Do not select a merely similar model.

Do not invent product information.

Return ONLY:

{{
    "selected_row_index": 1,
    "confidence": 0.99
}}

If there is no reliable match:

{{
    "selected_row_index": null,
    "confidence": 0.0
}}
"""

    response = groq_client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a deterministic "
                    "product-row selector. "
                    "Return only valid JSON."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        temperature=0,
        response_format={
            "type": "json_object",
        },
    )

    content = (
        response
        .choices[0]
        .message
        .content
    )

    if not content:

        raise ValueError(
            "Groq returned an empty response "
            "while selecting the product row."
        )

    try:

        result = json.loads(
            content
        )

    except json.JSONDecodeError as exc:

        raise ValueError(
            "Groq returned invalid JSON while "
            "selecting the product row."
        ) from exc

    selected_row_index = result.get(
        "selected_row_index"
    )

    try:

        confidence = float(
            result.get(
                "confidence",
                0,
            )
        )

    except (
        TypeError,
        ValueError,
    ):

        confidence = 0.0

    if selected_row_index is None:

        raise ValueError(
            f"Groq could not reliably identify "
            f"item code '{requested_item_code}'."
        )

    try:

        selected_row_index = int(
            selected_row_index
        )

    except (
        TypeError,
        ValueError,
    ) as exc:

        raise ValueError(
            "Groq returned an invalid row index."
        ) from exc

    if selected_row_index < 1:

        raise ValueError(
            "Groq selected the table header "
            "or an invalid row."
        )

    if (
        selected_row_index
        >= len(table.rows)
    ):

        raise ValueError(
            "Groq returned an out-of-range "
            "row index."
        )

    if confidence < 0.90:

        raise ValueError(
            f"Groq match confidence too low "
            f"for '{requested_item_code}': "
            f"{confidence}"
        )

    # --------------------------------------------------------
    # Final safety check.
    #
    # Even if Groq says a row is correct, make sure the row
    # actually contains enough evidence of the requested code.
    # --------------------------------------------------------

    selected_row = table.rows[
        selected_row_index
    ]

    selected_score = row_match_score(
        selected_row,
        requested_item_code,
    )

    if selected_score < 50:

        raise ValueError(
            f"Groq selected row "
            f"{selected_row_index}, but the row "
            f"does not contain sufficient evidence "
            f"for item '{requested_item_code}'."
        )

    return selected_row_index


# ============================================================
# Table Filtering
# ============================================================

def deepcopy_table_with_selected_row(
    source_table,
    selected_row_index,
):
    """
    Deep-copy a complete Word table while retaining:

        row 0 = header
        selected_row_index = requested product

    All other rows are removed.

    The original table remains untouched.
    """

    new_table_xml = deepcopy(
        source_table._tbl
    )

    copied_rows = new_table_xml.findall(
        qn("w:tr")
    )

    if not copied_rows:

        raise ValueError(
            "Source price-list table contains "
            "no rows."
        )

    if (
        selected_row_index < 1
        or selected_row_index >= len(
            copied_rows
        )
    ):

        raise ValueError(
            "Selected product row is outside "
            "the copied table."
        )

    rows_to_remove = []

    for index, row_xml in enumerate(
        copied_rows
    ):

        if index == 0:
            continue

        if index == selected_row_index:
            continue

        rows_to_remove.append(
            row_xml
        )

    for row_xml in rows_to_remove:

        new_table_xml.remove(
            row_xml
        )

    return new_table_xml


# ============================================================
# Complete Product Table Extraction
# ============================================================

def extract_product_item_table(
    source_doc,
    product_group,
    item_code,
    available_groups,
):
    """
    Locate the correct product-group table, identify the
    requested item row, and return a filtered table XML.

    Final table contains:

        1. Header
        2. Requested product row
    """

    tables = find_product_group_tables(
        source_doc=source_doc,
        product_group=product_group,
        available_groups=available_groups,
    )

    if not tables:

        raise ValueError(
            f"No price-list tables found for "
            f"product group '{product_group}'."
        )

    # --------------------------------------------------------
    # Identify the table containing the model.
    # --------------------------------------------------------

    table = find_table_containing_item(
        tables=tables,
        item_code=item_code,
    )

    # --------------------------------------------------------
    # Identify the exact product row.
    # --------------------------------------------------------

    selected_row_index = (
        find_table_row_with_groq(
            table=table,
            requested_item_code=item_code,
        )
    )

    # --------------------------------------------------------
    # Filter table.
    # --------------------------------------------------------

    return deepcopy_table_with_selected_row(
        source_table=table,
        selected_row_index=selected_row_index,
    )


# ============================================================
# Text Replacement
# ============================================================
def set_run_normal_text_font(
    run,
    font_name="Aptos",
):
    """
    Force a run to use a normal Unicode text font.

    This prevents Symbol/Wingdings/Webdings-style template
    formatting from turning normal replacement text into
    unrelated glyphs.
    """

    run.font.name = font_name

    rPr = run._r.get_or_add_rPr()

    rFonts = rPr.find(qn("w:rFonts"))

    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)

    rFonts.set(
        qn("w:ascii"),
        font_name,
    )

    rFonts.set(
        qn("w:hAnsi"),
        font_name,
    )

    rFonts.set(
        qn("w:cs"),
        font_name,
    )

    rFonts.set(
        qn("w:eastAsia"),
        font_name,
    )


def replace_text_in_paragraph(
    paragraph,
    replacements,
):
    """
    Replace placeholders inside a paragraph.

    Supports placeholders contained in one run or split
    across multiple runs.
    """

    for placeholder, replacement in (
        replacements.items()
    ):

        replacement = str(
            replacement or ""
        )

        full_text = paragraph.text

        if placeholder not in full_text:
            continue

        # ----------------------------------------------------
        # Placeholder in a single run
        # ----------------------------------------------------

        replaced = False

        for run in paragraph.runs:

            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement,)

                set_run_normal_text_font(run)
                replaced = True
                break

        if replaced:
            continue

        # ----------------------------------------------------
        # Placeholder split across runs
        # ----------------------------------------------------

        combined = "".join(
            run.text
            for run in paragraph.runs
        )

        if placeholder not in combined:
            continue

        new_text = combined.replace(
            placeholder,
            replacement,
        )

        if paragraph.runs:

            first_run = paragraph.runs[0]

            first_run.text = new_text

            set_run_normal_text_font(first_run)

            for run in paragraph.runs[1:]:

                run.text = ""


def replace_text_in_table(
    table,
    replacements,
):
    """
    Replace placeholders recursively inside Word tables.
    """

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                replace_text_in_paragraph(
                    paragraph,
                    replacements,
                )

            for nested_table in cell.tables:

                replace_text_in_table(
                    nested_table,
                    replacements,
                )


def replace_text_in_xml(
    element,
    replacements,
):
    """
    Replace placeholders in raw Word XML while also forcing
    replaced text to use a normal Unicode font.
    """

    for node in element.iter():

        if node.tag != qn("w:t"):
            continue

        if not node.text:
            continue

        for placeholder, replacement in replacements.items():

            if placeholder not in node.text:
                continue

            node.text = node.text.replace(
                placeholder,
                str(replacement or ""),
            )

            # Find the parent <w:r>
            run_element = node.getparent()

            if (
                run_element is not None
                and run_element.tag == qn("w:r")
            ):

                rPr = run_element.find(
                    qn("w:rPr")
                )

                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    run_element.insert(0, rPr)

                rFonts = rPr.find(
                    qn("w:rFonts")
                )

                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)

                for font_attribute in (
                    "ascii",
                    "hAnsi",
                    "cs",
                    "eastAsia",
                ):
                    rFonts.set(
                        qn(f"w:{font_attribute}"),
                        "Aptos",
                    )
                    

def replace_text_in_document(
    doc,
    replacements,
):
    """
    Replace placeholders throughout the complete quotation:

        - body paragraphs
        - body tables
        - headers
        - header tables
        - footers
        - footer tables
        - raw body XML
        - raw header/footer XML
    """

    # --------------------------------------------------------
    # Body paragraphs
    # --------------------------------------------------------

    for paragraph in doc.paragraphs:

        replace_text_in_paragraph(
            paragraph,
            replacements,
        )

    # --------------------------------------------------------
    # Body tables
    # --------------------------------------------------------

    for table in doc.tables:

        replace_text_in_table(
            table,
            replacements,
        )

    # --------------------------------------------------------
    # Headers / Footers
    # --------------------------------------------------------

    for section in doc.sections:

        for paragraph in (
            section.header.paragraphs
        ):

            replace_text_in_paragraph(
                paragraph,
                replacements,
            )

        for table in (
            section.header.tables
        ):

            replace_text_in_table(
                table,
                replacements,
            )

        for paragraph in (
            section.footer.paragraphs
        ):

            replace_text_in_paragraph(
                paragraph,
                replacements,
            )

        for table in (
            section.footer.tables
        ):

            replace_text_in_table(
                table,
                replacements,
            )

    # --------------------------------------------------------
    # Raw body XML
    # --------------------------------------------------------

    replace_text_in_xml(
        doc.element.body,
        replacements,
    )

    # --------------------------------------------------------
    # Raw header/footer XML
    # --------------------------------------------------------

    for section in doc.sections:

        replace_text_in_xml(
            section.header._element,
            replacements,
        )

        replace_text_in_xml(
            section.footer._element,
            replacements,
        )

def highlight_text_in_document(
    doc,
    search_text,
    color="00FFFF",
):
    """
    Highlight exact text throughout the document.

    Works across:
        - body paragraphs
        - body tables
        - headers
        - header tables
        - footers
        - footer tables

    color:
        Word highlight color value.
        00FFFF = turquoise/cyan.
    """

    def highlight_paragraph(paragraph):
        for run in paragraph.runs:
            if not run.text:
                continue

            if search_text in run.text:
                run.font.highlight_color = color

    # Body paragraphs
    for paragraph in doc.paragraphs:
        highlight_paragraph(paragraph)

    # Body tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    highlight_paragraph(paragraph)

                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                highlight_paragraph(paragraph)

    # Headers / footers
    for section in doc.sections:

        for paragraph in section.header.paragraphs:
            highlight_paragraph(paragraph)

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        highlight_paragraph(paragraph)

        for paragraph in section.footer.paragraphs:
            highlight_paragraph(paragraph)

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        highlight_paragraph(paragraph)


def highlight_product_group_headers(
    doc,
    product_group,
    color="00FFFF",
):
    """
    Highlight the product-group heading wherever it appears.
    """

    normalized_target = normalize_heading(
        product_group
    )

    def process_paragraph(paragraph):
        if normalize_heading(
            paragraph.text
        ) != normalized_target:
            return

        for run in paragraph.runs:
            run.font.highlight_color = color

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

# ============================================================
# Product Table Insertion
# ============================================================

def replace_product_group_with_content(
    target_doc,
    source_blocks,
    item_code,
    placeholder="{product_group}",
):
    """
    Replace {product_group} with copied paragraphs/tables.

    Returns the exact inserted table containing the requested item code.
    """

    body = target_doc.element.body

    for child in list(body):

        if not child.tag.endswith("}p"):
            continue

        paragraph = Paragraph(
            child,
            target_doc,
        )

        if placeholder not in paragraph.text:
            continue

        insert_index = list(body).index(child)

        # Remove placeholder paragraph
        body.remove(child)

        inserted_tables = []

        # Insert copied product-group content
        for block_xml in source_blocks:

            body.insert(
                insert_index,
                block_xml,
            )

            # Track inserted tables
            if block_xml.tag == qn("w:tbl"):

                inserted_tables.append(
                    Table(
                        block_xml,
                        target_doc,
                    )
                )

            insert_index += 1

        # ------------------------------------------------
        # Find the EXACT product table containing item code
        # ------------------------------------------------

        for table in inserted_tables:

            for row_index, row in enumerate(table.rows):

                if row_index == 0:
                    continue

                if row_contains_item_code(
                    row,
                    item_code,
                ):
                    return table

        # ------------------------------------------------
        # Fallback:
        # Find a 2-row table containing the product code
        # ------------------------------------------------

        for table in inserted_tables:

            if len(table.rows) == 2:

                if any(
                    row_contains_item_code(
                        row,
                        item_code,
                    )
                    for row in table.rows[1:]
                ):
                    return table

        raise ValueError(
            f"Product table containing item "
            f"'{item_code}' could not be identified "
            f"after insertion."
        )

    raise ValueError(
        f"Placeholder '{placeholder}' "
        f"not found in quotation template."
    )

def format_inr(value):
    value = money(value)

    integer_part = int(value)
    decimal_part = int(
        (value - integer_part) * 100
    )

    s = str(integer_part)

    if len(s) > 3:
        last_three = s[-3:]
        remaining = s[:-3]

        groups = []

        while len(remaining) > 2:
            groups.insert(0, remaining[-2:])
            remaining = remaining[:-2]

        if remaining:
            groups.insert(0, remaining)

        formatted = ",".join(groups + [last_three])

    else:
        formatted = s

    if decimal_part:
        return f"{formatted}.{decimal_part:02d}"

    return formatted

def create_quotation_calculation_table(
    target_doc,
    financials,
    packing_mode,
    freight_mode,
):
    """
    Create the quotation calculation section directly
    below the product-details table.

    No calculation placeholders are required.
    """

    table = target_doc.add_table(
        rows=0,
        cols=2,
    )

    table.style = "Table Grid"

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def add_row(label, value, bold=False):
        row = table.add_row()

        row.cells[0].text = label
        row.cells[1].text = value

        row.cells[0].vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        row.cells[1].vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"

                    if bold:
                        run.bold = True

        return row

    # ----------------------------------------------------
    # Packing
    # ----------------------------------------------------

    if packing_mode == "ACTUAL":
        packing_text = (
            f"₹ {format_inr(financials['packing_amount'])}"
        )
    else:
        packing_text = "INCLUSIVE"

    add_row(
        "(+) Add : Packing Charges",
        f"→ {packing_text}",
    )

    # ----------------------------------------------------
    # Freight
    # ----------------------------------------------------

    if freight_mode == "ACTUAL":
        freight_text = (
            f"₹ {format_inr(financials['freight_amount'])}"
        )
    else:
        freight_text = "INCLUSIVE"

    add_row(
        "(+) Add : Freight Charges",
        f"→ {freight_text}",
    )

    # ----------------------------------------------------
    # Gross total / taxable value
    # ----------------------------------------------------

    add_row(
        "Gross Total",
        f"→ ₹ {format_inr(financials['taxable_value'])}",
        bold=True,
    )

    # ----------------------------------------------------
    # GST
    # ----------------------------------------------------

    add_row(
        f"(+) Add : GST {format_inr(financials['tax_rate'])}%",
        f"→ ₹ {format_inr(financials['tax_amount'])}",
    )

    # ----------------------------------------------------
    # Grand total
    # ----------------------------------------------------

    add_row(
        "Gross Total Amount",
        f"→ ₹ {format_inr(financials['grand_total'])}",
        bold=True,
    )

    return table

def add_financial_rows_to_product_table(product_table, financials, packing_mode, freight_mode,):
    """
    Append financial rows directly to the exact product table.

    The product table may have multiple columns.
    Financial rows are merged into:

        [ Description spanning most columns ] [ Amount ]
    """

    if product_table is None:
        raise ValueError("Product table is required for adding financial rows.")

    column_count = len(product_table.columns)

    if column_count < 2:
        raise ValueError("Product table must contain at least 2 columns.")

    # ====================================================
    # Helper
    # ====================================================

    def add_financial_row(label, value, bold=False,):
        row = product_table.add_row()

        # ------------------------------------------------
        # Merge all columns except the final column.
        #
        # Example for an 8-column product table:
        #
        # |              LABEL               | VALUE |
        # | col0 col1 col2 col3 col4 col5 c6 | col7  |
        # ------------------------------------------------

        label_cell = row.cells[0]

        if column_count > 2:

            for index in range(1, column_count - 1,):
                label_cell = label_cell.merge(row.cells[index])

        value_cell = row.cells[-1]

        # ------------------------------------------------
        # Set label
        # ------------------------------------------------

        label_cell.text = str(label)

        label_cell.vertical_alignment = (WD_CELL_VERTICAL_ALIGNMENT.CENTER)

        # ------------------------------------------------
        # Set value
        # ------------------------------------------------

        value_cell.text = str(value)

        value_cell.vertical_alignment = (WD_CELL_VERTICAL_ALIGNMENT.CENTER)

        # ------------------------------------------------
        # Font formatting
        # ------------------------------------------------

        for cell in (label_cell, value_cell,):

            for paragraph in cell.paragraphs:

                for run in paragraph.runs:

                    run.font.name = "Aptos"
                    run.bold = bold

        return row

    # ====================================================
    # Packing
    # ====================================================

    if packing_mode == "ACTUAL":

        packing_text = (f"₹ {format_inr(financials['packing_amount'])}")

    else:

        packing_text = "INCLUSIVE"

    add_financial_row("(+) Add : Packing Charges", f"→ {packing_text}",)

    # ====================================================
    # Freight
    # ====================================================

    if freight_mode == "ACTUAL":

        freight_text = (f"₹ {format_inr(financials['freight_amount'])}")

    else:

        freight_text = "INCLUSIVE"

    add_financial_row("(+) Add : Freight Charges", f"→ {freight_text}",)

    # ====================================================
    # Gross Total / Taxable Value
    # ====================================================

    add_financial_row("Gross Total", (f"→ ₹ " f"{format_inr(financials['taxable_value'])}"), bold=True,)

    # ====================================================
    # GST
    # ====================================================

    add_financial_row((f"(+) Add : GST " f"{format_inr(financials['tax_rate'])}%"), ( f"→ ₹ " f"{format_inr(financials['tax_amount'])}"),)

    # ====================================================
    # Grand Total
    # ====================================================

    add_financial_row(
        "Gross Total Amount",
        (
            f"→ ₹ "
            f"{format_inr(financials['grand_total'])}"
        ),
        bold=True,
    )

    return product_table

# ============================================================
# Special Model Table
# ============================================================

def get_special_row_values(
    row_data,
):
    """
    Support dictionary rows and objects with a `.values`
    attribute.
    """

    if isinstance(
        row_data,
        dict,
    ):

        return row_data.get(
            "values",
            [],
        )

    values = getattr(
        row_data,
        "values",
        None,
    )

    if values is not None:
        return values

    return []


def create_special_model_table(
    target_doc,
    columns,
    rows,
):
    """
    Create a dynamic special-model configuration table.

    Replaces:

        {special_model_table}
    """

    columns = [
        str(column or "").strip()
        for column in (
            columns or []
        )
    ]

    rows = rows or []

    columns = [
        column
        if column
        else f"Column {index + 1}"
        for index, column in enumerate(
            columns
        )
    ]

    if not columns or not rows:
        return None

    # --------------------------------------------------------
    # Locate placeholder
    # --------------------------------------------------------

    placeholder_paragraph = None

    for paragraph in (
        target_doc.paragraphs
    ):

        if (
            "{special_model_table}"
            in paragraph.text
        ):

            placeholder_paragraph = (
                paragraph
            )

            break

    # --------------------------------------------------------
    # Search body tables
    # --------------------------------------------------------

    if placeholder_paragraph is None:

        for table in target_doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in (
                        cell.paragraphs
                    ):

                        if (
                            "{special_model_table}"
                            in paragraph.text
                        ):

                            placeholder_paragraph = (
                                paragraph
                            )

                            break

                    if placeholder_paragraph:
                        break

                if placeholder_paragraph:
                    break

            if placeholder_paragraph:
                break

    if placeholder_paragraph is None:

        raise ValueError(
            "Placeholder "
            "'{special_model_table}' "
            "not found."
        )

    # --------------------------------------------------------
    # Create table
    # --------------------------------------------------------

    table = target_doc.add_table(
        rows=1,
        cols=len(columns),
    )

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    table.style = "Table Grid"

    # --------------------------------------------------------
    # Header
    # --------------------------------------------------------

    header_cells = (
        table.rows[0].cells
    )

    for index, column in enumerate(
        columns
    ):

        cell = header_cells[index]

        cell.text = column

        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        for paragraph in (
            cell.paragraphs
        ):

            for run in paragraph.runs:

                run.bold = True

    # --------------------------------------------------------
    # Data rows
    # --------------------------------------------------------

    for row_data in rows:

        values = get_special_row_values(
            row_data
        )

        values = [
            str(value or "")
            for value in values
        ]

        if len(values) < len(columns):

            values.extend(
                [""] * (
                    len(columns)
                    - len(values)
                )
            )

        elif len(values) > len(columns):

            values = values[
                :len(columns)
            ]

        cells = (
            table
            .add_row()
            .cells
        )

        for index, value in enumerate(
            values
        ):

            cells[index].text = value

            cells[index].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

    # --------------------------------------------------------
    # Insert at placeholder position
    # --------------------------------------------------------

    placeholder_p = (
        placeholder_paragraph._p
    )

    table_element = table._tbl

    placeholder_p.addnext(
        table_element
    )

    parent = (
        placeholder_p.getparent()
    )

    if parent is not None:

        parent.remove(
            placeholder_p
        )

    return table

def deepcopy_product_group_blocks(
    source_doc,
    product_group,
    available_groups,
):
    """
    Deep-copy every Word block belonging to the requested
    product group.

    Includes:
        - paragraphs
        - headings
        - specification text
        - tables

    The original source document is never modified.

    Returns:
        list[deepcopied XML elements]
    """

    (
        group_start,
        group_end,
        blocks,
    ) = find_product_group_bounds(
        source_doc=source_doc,
        product_group=product_group,
        available_groups=available_groups,
    )

    selected_blocks = blocks[
        group_start:group_end
    ]

    if not selected_blocks:
        raise ValueError(
            f"No content found for product group "
            f"'{product_group}'."
        )

    copied_blocks = []

    for block in selected_blocks:

        if isinstance(block, Paragraph):

            copied_blocks.append(
                deepcopy(block._p)
            )

        elif isinstance(block, Table):

            copied_blocks.append(
                deepcopy(block._tbl)
            )

    return copied_blocks

def filter_table_xml_to_selected_row(
    table_xml,
    selected_row_index,
):
    """
    Filter a deep-copied Word table XML.

    Keeps:
        row 0 = header
        selected_row_index = requested product

    Removes all other rows.

    The source document remains untouched.
    """

    rows = table_xml.findall(
        qn("w:tr")
    )

    if not rows:
        raise ValueError(
            "Copied price-list table contains no rows."
        )

    if (
        selected_row_index < 1
        or selected_row_index >= len(rows)
    ):
        raise ValueError(
            "Selected product row is outside "
            "the copied table."
        )

    for index, row_xml in enumerate(
        list(rows)
    ):

        if index == 0:
            continue

        if index == selected_row_index:
            continue

        table_xml.remove(
            row_xml
        )

    return table_xml

def extract_product_group_content(
    source_doc,
    product_group,
    item_code,
    available_groups,
):
    """
    Extract the complete requested product-group section.

    Workflow:

        1. Locate product-group boundaries.
        2. Deep-copy all blocks in that group.
        3. Find the source table containing the requested model.
        4. Ask Groq for the exact matching row when required.
        5. Filter the copied table.
        6. Return the complete copied block sequence.

    The returned content contains:
        - product-group heading
        - specifications
        - descriptive paragraphs
        - filtered product table
    """

    (
        group_start,
        group_end,
        blocks,
    ) = find_product_group_bounds(
        source_doc=source_doc,
        product_group=product_group,
        available_groups=available_groups,
    )

    group_blocks = blocks[
        group_start:group_end
    ]

    if not group_blocks:
        raise ValueError(
            f"No content found for product group "
            f"'{product_group}'."
        )

    # --------------------------------------------------------
    # Find source tables in this product group
    # --------------------------------------------------------

    tables = [
        block
        for block in group_blocks
        if isinstance(block, Table)
    ]

    if not tables:
        raise ValueError(
            f"No price-list tables found for "
            f"product group '{product_group}'."
        )

    # --------------------------------------------------------
    # Identify correct table
    # --------------------------------------------------------

    source_table = find_table_containing_item(
        tables=tables,
        item_code=item_code,
    )

    # --------------------------------------------------------
    # Identify exact requested row
    # --------------------------------------------------------

    selected_row_index = (
        find_table_row_with_groq(
            table=source_table,
            requested_item_code=item_code,
        )
    )

    # --------------------------------------------------------
    # Deep-copy EVERYTHING in product group
    # --------------------------------------------------------

    copied_blocks = []

    for block in group_blocks:

        if isinstance(block, Paragraph):

            copied_blocks.append(
                deepcopy(block._p)
            )

        elif isinstance(block, Table):

            copied_table_xml = deepcopy(
                block._tbl
            )

            # ----------------------------------------------
            # This is the table containing the requested item
            # ----------------------------------------------

            if block is source_table:

                copied_table_xml = (
                    filter_table_xml_to_selected_row(
                        table_xml=copied_table_xml,
                        selected_row_index=(
                            selected_row_index
                        ),
                    )
                )

            copied_blocks.append(
                copied_table_xml
            )

    return copied_blocks

def replace_special_model_placeholder(
    target_doc,
    special_model,
    special_columns=None,
    special_rows=None,
):
    """
    Render the special-model table when enabled.

    Otherwise remove its placeholder.
    """

    if special_model:

        create_special_model_table(
            target_doc=target_doc,
            columns=special_columns or [],
            rows=special_rows or [],
        )

        return

    replace_text_in_document(
        target_doc,
        {"{special_model_table}": ""},
    )
    

# ============================================================
# Dealer Row
# ============================================================

def add_dealer_row(
    target_doc,
    dealer,
):
    """
    Insert the dealer row into the quotation information table.

    Existing row index 2 is used as the formatting template.
    """

    if not dealer:
        return

    if not target_doc.tables:

        raise ValueError(
            "Quotation template does not "
            "contain a dealer information table."
        )

    table = target_doc.tables[0]

    if len(table.rows) < 3:

        raise ValueError(
            "Quotation information table "
            "does not contain the expected "
            "source row."
        )

    source_row = table.rows[2]

    new_row = deepcopy(
        source_row._tr
    )

    table._tbl.insert(
        1,
        new_row,
    )

    row = table.rows[1]

    row.cells[0].text = "DEALER"

    row.cells[1].text = (
        "Dealer quotation applicable."
    )

    for index in range(
        2,
        len(row.cells),
    ):

        row.cells[index].text = ""


# ============================================================
# Main Quotation Generator
# ============================================================

def generate_qoute_document(
    request,
    authenticated_user=None,
    sales_user=None,
):

    if not authenticated_user:

        raise PermissionError("Authentication is required " "to generate quotations.")

    allowed_roles = {"Sales Representative", "Admin", "Chief Full Stack Developer",}

    user_role = authenticated_user.get("role")

    if user_role not in allowed_roles:

        raise PermissionError("Only sales team can generate " "quotations.")

    user_email = authenticated_user.get("email")

    if not user_email:

        raise PermissionError("Authenticated user does not " "have an email address.")

    # ========================================================
    # Sales User
    # ========================================================

    sales_user = (
        EDBR.get_user_business_contact(
            email=user_email,
            role=user_role,
        )
    )

    if not sales_user:

        raise ValueError(
            "Could not find the business "
            "contact information for the "
            "authenticated sales user."
        )

    # ========================================================
    # Request Validation
    # ========================================================

    if not request:

        raise ValueError(
            "Quotation request is required."
        )

    if not request.product_group:

        raise ValueError(
            "Product group is required."
        )

    if not request.item_code:

        raise ValueError(
            "Item code is required."
        )

    if not request.qoute_number:

        raise ValueError(
            "Quotation number is required."
        )

    if not request.client_company:

        raise ValueError(
            "Client company is required."
        )
    if request.base_model_price is None:
        raise ValueError(
            "Base model price is required."
        )

    if request.base_model_price < 0:
        raise ValueError(
            "Base model price cannot be negative."
        )

    if request.packing_mode not in {
        "INCLUSIVE",
        "ACTUAL",
    }:
        raise ValueError(
            "Invalid packing mode."
        )

    if request.freight_mode not in {
        "INCLUSIVE",
        "ACTUAL",
    }:
        raise ValueError(
            "Invalid freight mode."
        )

    if request.packing_mode == "ACTUAL":
        if request.packing_amount is None:
            raise ValueError(
                "Packing amount is required when packing mode is ACTUAL."
            )

    if request.freight_mode == "ACTUAL":
        if request.freight_amount is None:
            raise ValueError(
                "Freight amount is required when freight mode is ACTUAL."
            )

    quote_date = request.date_input

    if quote_date is None:

        raise ValueError(
            "Quotation date is required."
        )

    financials = calculate_quotation_financials(request)

    # ========================================================
    # File Validation
    # ========================================================

    if not TEMPLATE_PATH.exists():

        raise FileNotFoundError(
            f"Quotation template not found: "
            f"{TEMPLATE_PATH}"
        )

    if not PRICE_LIST_PATH.exists():

        raise FileNotFoundError(
            f"Price list not found: "
            f"{PRICE_LIST_PATH}"
        )

    # ========================================================
    # Load Documents
    # ========================================================

    source_doc = Document(
        PRICE_LIST_PATH
    )

    target_doc = Document(
        TEMPLATE_PATH
    )

    # ========================================================
    # Product Groups
    # ========================================================

    available_products = (
        EDBR.get_item_groups()
    )

    if not available_products:

        raise ValueError(
            "No product groups are available "
            "in the database."
        )

    # ========================================================
    # Extract Filtered Price Table
    # ========================================================

    product_group_content = (
    extract_product_group_content(
        source_doc=source_doc,
        product_group=request.product_group,
        item_code=request.item_code,
        available_groups=available_products,
    )
    )

    # ========================================================
    # Insert Product Table
    # ========================================================

    product_table = replace_product_group_with_content(target_doc=target_doc, source_blocks=product_group_content, item_code=request.item_code,)

    add_financial_rows_to_product_table(product_table=product_table, financials=financials, packing_mode=request.packing_mode, freight_mode=request.freight_mode,)
    # ========================================================
    # Dealer
    # ========================================================

    if request.dealer:

        add_dealer_row(
            target_doc=target_doc,
            dealer=request.dealer,
        )

    # ========================================================
    # Special Model
    # ========================================================

    replace_special_model_placeholder(
        target_doc=target_doc,
        special_model=bool(
            request.special_model
        ),
        special_columns=(
            request.special_columns
        ),
        special_rows=(
            request.special_rows
        ),
    )

    # ========================================================
    # Normal Placeholder Values
    # ========================================================

    replacements = {
        "{client_company}":
            request.client_company,

        "{client_address_line_1}":
            request.client_address_line1,

        "{client_city-client_postal_code}":
            (
                f"{request.client_city} - "
                f"{request.client_postal_code}"
            ),

        "{client_email}":
            request.client_email,

        "{buyer_name}":
            request.buyer_name,

        "{buyer_phone_number}":
            request.buyer_phone_number,

        "{date_input}":
            quote_date.strftime(
                "%d/%m/%Y"
            ),

        "{current_date}":
            quote_date.strftime(
                "%d/%m/%Y"
            ),

        "{quote_num}":
            str(
                request.qoute_number
            ),

        "{sales_user_name}":
            sales_user.get(
                "name",
                "",
            ),

        "{business_phone}":
            sales_user.get(
                "business_phone",
                "",
            ) or "",

        "{sales_user_email}":
            sales_user.get(
                "email",
                "",
            ),

        "{supply}":
            request.supply,

        "{installation}":
            request.installation,

        "{freight}":
            request.freight,
    }

    replace_text_in_document(target_doc, replacements,)

    highlight_product_group_headers(
    doc=target_doc,
    product_group=request.product_group,
    color=WD_COLOR_INDEX.TURQUOISE,
    )

    # Highlight the sales-user contact XML text.
    highlight_text_in_document(doc=target_doc, search_text=(f"{sales_user.get('name', '')}(:{sales_user.get('business_phone', '') or ''})"),
        color=WD_COLOR_INDEX.TURQUOISE,
    )

    # ========================================================
    # Output Directory
    # ========================================================

    OUTPUT_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    # ========================================================
    # Output Filename
    # ========================================================

    safe_quote_number = safe_filename(
        request.qoute_number
    )

    safe_company = safe_filename(
        request.client_company
    )

    filename = (
        f"Tempo_Qoute_"
        f"{safe_quote_number}_"
        f"{safe_company}.docx"
    )

    output_path = (
        OUTPUT_DIR
        / filename
    )

    # ========================================================
    # Save
    # ========================================================

    target_doc.save(
        output_path
    )

    # ========================================================
    # Return
    # ========================================================

    return (
        output_path,
        sales_user,
    )