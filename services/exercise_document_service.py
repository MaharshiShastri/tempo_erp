from datetime import date
from pathlib import Path
from io import BytesIO
from copy import deepcopy

from docx import Document


TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent
    / "services"
    / "exercise_tracking_template.docx"
)


def _replace_text_in_paragraph(
    paragraph,
    replacements: dict[str, str],
):
    """
    Replace placeholders even when Word has split the placeholder
    across multiple runs.

    Example:
        "{exercise" + "_name}"
    """

    for old, new in replacements.items():

        while old in paragraph.text:

            runs = paragraph.runs

            if not runs:
                break

            full_text = "".join(
                run.text or ""
                for run in runs
            )

            start = full_text.find(old)

            if start == -1:
                break

            end = start + len(old)

            # -----------------------------------------------------
            # Find the run containing the start/end of placeholder
            # -----------------------------------------------------

            start_run_index = None
            end_run_index = None

            start_offset = None
            end_offset = None

            cursor = 0

            for index, run in enumerate(runs):

                run_text = run.text or ""

                run_start = cursor
                run_end = cursor + len(run_text)

                if (
                    start_run_index is None
                    and run_start <= start < run_end
                ):
                    start_run_index = index
                    start_offset = start - run_start

                if (
                    run_start < end <= run_end
                ):
                    end_run_index = index
                    end_offset = end - run_start
                    break

                cursor = run_end

            # -----------------------------------------------------
            # Fallback for empty/special runs
            # -----------------------------------------------------

            if (
                start_run_index is None
                or end_run_index is None
            ):
                break

            start_run = runs[start_run_index]
            end_run = runs[end_run_index]

            start_text = start_run.text or ""
            end_text = end_run.text or ""

            # Same run
            if start_run_index == end_run_index:

                start_run.text = (
                    start_text[:start_offset]
                    + str(new)
                    + start_text[end_offset:]
                )

            else:

                # Keep everything before placeholder
                prefix = start_text[:start_offset]

                # Keep everything after placeholder
                suffix = end_text[end_offset:]

                start_run.text = (
                    prefix
                    + str(new)
                )

                # Clear all runs between start/end
                for index in range(
                    start_run_index + 1,
                    end_run_index,
                ):
                    runs[index].text = ""

                end_run.text = suffix

def _replace_text_in_table(
    table,
    replacements: dict[str, str],
):
    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(
                    paragraph,
                    replacements,
                )

            for nested_table in cell.tables:
                _replace_text_in_table(
                    nested_table,
                    replacements,
                )


def _clone_table_row(
    table,
    source_row,
):
    """
    Clone a Word table row while preserving
    its formatting, borders, widths, etc.
    """

    new_tr = deepcopy(source_row._tr)

    table._tbl.append(new_tr)

    return table.rows[-1]


def _clear_cell(cell):
    """
    Remove all existing content from a cell
    and leave one empty paragraph.
    """

    tc = cell._tc

    for child in list(tc):
        if child.tag.endswith("}tcPr"):
            continue

        tc.remove(child)

    cell.add_paragraph()


def _set_cell_text(cell, text):
    """
    Replace cell content with plain text while
    preserving the cell itself and its formatting.
    """

    for paragraph in cell.paragraphs:

        for run in paragraph.runs:
            run.text = ""

    if cell.paragraphs:

        cell.paragraphs[0].text = str(text)

    else:

        cell.text = str(text)


def _populate_name_rows(
    table,
    person_names: list[str],
):
    """
    First row = header.
    Second row = template data row.

    Duplicate the template row for every person.
    """

    if len(table.rows) < 2:
        raise ValueError(
            "Exercise template must contain a header row "
            "and at least one data row."
        )

    template_row = table.rows[1]

    # Remove all existing rows after the header.
    while len(table.rows) > 1:
        table._tbl.remove(
            table.rows[1]._tr
        )

    for index, person_name in enumerate(
        person_names
    ):

        if index == 0:
            row = table.rows[1] if len(table.rows) > 1 else _clone_table_row(
                table,
                template_row,
            )
        else:
            row = _clone_table_row(
                table,
                template_row,
            )

        # Ensure the first cell contains ONLY this name.
        _set_cell_text(
            row.cells[0],
            person_name,
        )

        # Checkbox column.
        _set_cell_text(
            row.cells[1],
            "☐",
        )

        # Remaining fields blank.
        _set_cell_text(row.cells[2], "")
        _set_cell_text(row.cells[3], "")
        _set_cell_text(row.cells[4], "")


def generate_exercise_document(
    *,
    exercise_name: str,
    person_names: list[str],
) -> BytesIO:

    exercise_name = (exercise_name or "").strip()

    if not exercise_name:
        raise ValueError(
            "Exercise name is required."
        )

    cleaned_names = [
        name.strip()
        for name in person_names
        if name and name.strip()
    ]

    if not cleaned_names:
        raise ValueError(
            "At least one person is required."
        )

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            "Exercise tracking template was not found."
        )

    document = Document(
        str(TEMPLATE_PATH)
    )

    replacements = {
        "{exercise_name}": exercise_name,
        "{current_date}": date.today().strftime(
            "%d-%m-%Y"
        ),
    }

    # ---------------------------------------------------------
    # Replace placeholders in normal paragraphs
    # ---------------------------------------------------------

    for paragraph in document.paragraphs:
        _replace_text_in_paragraph(
            paragraph,
            replacements,
        )

    # ---------------------------------------------------------
    # First table = acknowledgement table
    # ---------------------------------------------------------

    if document.tables:

        exercise_table = document.tables[0]

        _populate_name_rows(
            exercise_table,
            cleaned_names,
        )

    # ---------------------------------------------------------
    # Replace placeholders inside all tables
    # ---------------------------------------------------------

    for table in document.tables:

        _replace_text_in_table(
            table,
            replacements,
        )

    # ---------------------------------------------------------
    # Also handle headers and footers, just in case
    # ---------------------------------------------------------

    for section in document.sections:

        for paragraph in section.header.paragraphs:
            _replace_text_in_paragraph(
                paragraph,
                replacements,
            )

        for table in section.header.tables:
            _replace_text_in_table(
                table,
                replacements,
            )

        for paragraph in section.footer.paragraphs:
            _replace_text_in_paragraph(
                paragraph,
                replacements,
            )

        for table in section.footer.tables:
            _replace_text_in_table(
                table,
                replacements,
            )

    # ---------------------------------------------------------
    # Save in memory
    # ---------------------------------------------------------

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output