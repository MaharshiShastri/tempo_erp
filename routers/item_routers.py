from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from schemas.items_schema import ItemMasterCreate, ItemMasterUpdate
from database.repository import EDBR
from security import verify_bearer_token
from .dependencies import check_department
import docx
import io

router = APIRouter(prefix="/api/v1/master/items", tags=["Item Master Subsystem"])

@router.get("/{item_code}")
def list_items(item_code: str, user_profile: dict=Depends(verify_bearer_token)):
    return EDBR.get_item(item_code)

@router.post("/create")
def create_item(payload: ItemMasterCreate, user_profile: dict=Depends(verify_bearer_token)):
    try:
        return EDBR.create_item(payload)
    
    except Exception as e:
        raise HTTPException(status_code=400, detail="Item Code already exists or data is invalid.")
    
@router.put("/{item_code}")
def update_item(item_code: str, payload: ItemMasterUpdate, user_profile=Depends(verify_bearer_token)):
    return EDBR.update_item(item_code, payload.dict(exclude_none=True))

@router.delete("/{item_code}")
def delete_item(item_code: str,user_profile=Depends(verify_bearer_token)):
    return EDBR.disable_item(item_code)

@router.post("/upload-word")
async def upload_word_catalog(file: UploadFile = File(...), user: dict = Depends(verify_bearer_token)):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    
    try:
        contents = await file.read()
        doc = docx.Document(io.BytesIO(contents))
        
        inserted_count = 0
        # Assuming the first table in the document holds the catalog data
        if not doc.tables:
            raise ValueError("No tables found in the Word document.")
            
        table = doc.tables[0]
        
        # Skip header row, iterate through data
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 4: # Assuming format: [Code, Name, Group, Rate]
                item_data = {
                    "item_code": cells[0],
                    "item_name": cells[1],
                    "item_group": cells[2],
                    "rate": float(cells[3]) if cells[3].replace('.','',1).isdigit() else 0.0,
                    "unit_measure": "NOS",
                    "additional_spec_text": "",
                    "hsn_code": "",
                    "revision_no": "0"
                }
                EDBR.create_item(item_data)
                inserted_count += 1
                
        return {"status": "success", "inserted": inserted_count}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Word document: {str(e)}")
    
"""@router.get("/search")
def search_items(q: str):
    with EDBR._get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("""
   #             SELECT item_code, item_name, hsn_code, rate
    #            FROM items_master
     #           WHERE is_active = TRUE
      #            AND (
       #                 item_code ILIKE %s OR
        #                item_name ILIKE %s
         #         )
          #      ORDER BY item_name
           #     LIMIT 10
""", (f"%{q}%", f"%{q}%"))

return cur.fetchall()"""